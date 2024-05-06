import tkinter
from tkinter import ttk
from tkinter import messagebox
from docx import Document
import datetime

def generate_invoice():
    customer_name = first_nameEntry.get()+" "+last_nameEntry.get()
    # Retrieve customer name from first_nameEntry
    # Get other invoice details from other input fields
    
    # Create a new Word document
    doc = Document()
    
    # Populate the document with invoice details
    doc.add_heading('Invoice', level=1)
    doc.add_paragraph(f'Customer Name: {customer_name}')
    # Add other details
    
    # Save the document
    doc.save('invoice.docx')
    
    messagebox.showinfo("Invoice Generated", "Invoice has been generated successfully!")

def reset():
    first_nameEntry.delete(0, tkinter.END)
    last_nameEntry.delete(0, tkinter.END)
    phoneEntry.delete(0, tkinter.END)
    clear_item()
    tree.delete(*tree.get_children())

def clear_item():
    quantitySpinbox.delete(0, tkinter.END)
    quantitySpinbox.insert(0, "1")  # used to Re-enter the erased value
    desEntry.delete(0, tkinter.END)
    priceEntry.delete(0, tkinter.END)
    priceEntry.insert(0, "0.0")

def add_data():
    quantity = int(quantitySpinbox.get())
    description = desEntry.get()
    price = float(priceEntry.get())
    total = quantity * price
    invoice_item = [quantity, description, price, total]
    
    tree.insert('', 0, values=invoice_item)
    clear_item()

window = tkinter.Tk()
window.title("Invoice Generator Form")

frame = tkinter.Frame(window)
frame.pack(padx=20, pady=25)

first_nameLabel = tkinter.Label(frame, text="First name")
first_nameLabel.grid(row=0, column=0)

last_nameLabel = tkinter.Label(frame, text="Last name")
last_nameLabel.grid(row=0, column=1)

first_nameEntry = tkinter.Entry(frame)
last_nameEntry = tkinter.Entry(frame)
first_nameEntry.grid(row=1, column=0)
last_nameEntry.grid(row=1, column=1)

phoneLabel = tkinter.Label(frame, text="Phone")
phoneLabel.grid(row=0, column=2)
phoneEntry = tkinter.Entry(frame)
phoneEntry.grid(row=1, column=2)

quantityLabel = tkinter.Label(frame, text="Quantity")
quantityLabel.grid(row=2, column=0)
quantitySpinbox = tkinter.Spinbox(frame, from_=1, to=100)
quantitySpinbox.grid(row=3, column=0)

description = tkinter.Label(frame, text="Description:")
description.grid(row=3, column=1)
desEntry = tkinter.Entry(frame)
desEntry.grid(row=3, column=2)

priceLabel = tkinter.Label(frame, text="Unit Price")
priceLabel.grid(row=4, column=0)
priceEntry = tkinter.Spinbox(frame, from_=0.5, to=500, increment=0.5)
priceEntry.grid(row=5, column=0)

addItem = tkinter.Button(frame, text="Add Item", command=add_data)
addItem.grid(row=6, column=3)

column = ('quantity', 'description', 'price', 'total')

tree = ttk.Treeview(frame, columns=column, show="headings")
tree.heading('quantity', text="Quantity")
tree.heading('description', text="Description")
tree.heading('price', text="Unit Price")
tree.heading('total', text="Total")

tree.grid(row=7, column=0, columnspan=3, padx=10)

generateLabel = tkinter.Button(frame, text="Generate Invoice", command=generate_invoice)
generateLabel.grid(row=8, column=1, columnspan=3, pady=10)

resetLabel = tkinter.Button(frame, text="Fresh Invoice", command=reset)
resetLabel.grid(row=8, column=0, pady=10)

window.mainloop()
